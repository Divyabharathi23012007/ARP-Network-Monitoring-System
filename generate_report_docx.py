import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import os

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def create_report():
    doc = Document()

    # Set standard 1-inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Styles
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x11, 0x18, 0x27)

    def add_main_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(14)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        return p

    def add_section_header(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(13)
        run.font.name = 'Times New Roman'
        run.underline = True
        run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        return p

    def add_sub_header(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(11.5)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
        return p

    def add_body_p(text, bold_prefix="", italic=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            r_prefix = p.add_run(bold_prefix)
            r_prefix.bold = True
            r_prefix.font.name = 'Times New Roman'
        r = p.add_run(text)
        r.italic = italic
        r.font.name = 'Times New Roman'
        return p

    # -------------------------------------------------------------
    # PAGE 1: TITLE & PO-PSO MAPPING TABLE
    # -------------------------------------------------------------
    add_main_title("ARP NETWORK MONITORING SYSTEM – DETECT ABNORMAL CHANGES IN SIMULATED ARP MAPPINGS")

    add_sub_header("PO-PSO Mapping Table")

    # Table 1: Matrix
    table_matrix = doc.add_table(rows=2, cols=15)
    table_matrix.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_matrix.autofit = False

    headers_m = ["PO1", "PO2", "PO3", "PO4", "PO5", "PO6", "PO7", "PO8", "PO9", "PO10", "PO11", "PO12", "PSO1", "PSO2", "PSO3"]
    values_m = ["3", "3", "3", "2", "3", "2", "1", "1", "2", "2", "2", "2", "3", "3", "3"]

    hdr_cells = table_matrix.rows[0].cells
    for i, title in enumerate(headers_m):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "F1F5F9")
        set_cell_margins(hdr_cells[i], 60, 60, 40, 40)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(8.5)
        p.runs[0].font.name = 'Times New Roman'

    val_cells = table_matrix.rows[1].cells
    for i, val in enumerate(values_m):
        val_cells[i].text = val
        set_cell_margins(val_cells[i], 60, 60, 40, 40)
        p = val_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.name = 'Times New Roman'

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # PO JUSTIFICATION TABLE
    add_sub_header("PO JUSTIFICATION")

    po_data = [
        ("PO1", "Engineering Knowledge", "3", "Applies core Computer Networking principles, Data Link layer mechanics (RFC 826), Ethernet framing, Java 17 backend concurrency, and React web engineering to build a real-time network defense system."),
        ("PO2", "Problem Analysis", "3", "Identifies and mathematically models security vulnerabilities arising from ARP's stateless architecture, analyzing ARP cache poisoning patterns, MAC flip-flop churn, and rate anomalies."),
        ("PO3", "Design/Development", "3", "Designs and implements a full-stack network monitoring solution with Spring Boot 3 REST APIs, live WebSocket event pipelines, in-memory switch simulation, and an interactive UI dashboard."),
        ("PO4", "Investigation", "2", "Investigates simulated packet flows, inspects Layer-2 anomalies using Dynamic ARP Inspection (DAI) baselines, and traces attack vectors using real-time packet decoders."),
        ("PO5", "Modern Tools", "3", "Leverages industry-standard modern engineering tools including Java 17, Spring Boot 3, Maven, Docker, React 18, WebSockets, HTML5 Canvas 60FPS rendering, Chart.js, and Tailwind CSS."),
        ("PO6", "Engineering & Society", "2", "Protects local area networks against Man-in-the-Middle (MITM) eavesdropping, credential theft, and communication interception in educational and enterprise environments."),
        ("PO7", "Environment", "1", "Promotes efficient compute and network bandwidth utilization through lightweight sliding-window rate tracking and memory-optimized virtual state machines."),
        ("PO8", "Ethics", "1", "Emphasizes cybersecurity defense, safe defensive simulation within virtual sandboxes, and responsible network vulnerability research without harming physical network infrastructure."),
        ("PO9", "Individual & Team Work", "2", "Modular codebase architecture allows collaborative extension across frontend telemetry visualization, backend simulation engines, and automated mitigation modules."),
        ("PO10", "Communication", "2", "Provides clear visual topology diagrams, real-time alert logs, deep packet breakdown modals, and exportable forensic audit reports for technical evaluation."),
        ("PO11", "Project Management", "2", "Demonstrates end-to-end software engineering lifecycle including requirement specification, architecture modeling, test-driven validation, and 1-click execution packaging."),
        ("PO12", "Lifelong Learning", "2", "Encourages continuous exploration of next-generation network protocols, IPv6 Neighbor Discovery Protocol (NDP), Cisco DAI switch security, and AI-driven intrusion prevention systems.")
    ]

    table_po = doc.add_table(rows=1, cols=4)
    table_po.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_po.autofit = False

    po_hdrs = table_po.rows[0].cells
    po_cols = ["PO", "Description", "Relevance", "Justification"]
    col_widths_po = [Inches(0.8), Inches(1.8), Inches(0.9), Inches(3.0)]

    for i, h in enumerate(po_cols):
        po_hdrs[i].text = h
        po_hdrs[i].width = col_widths_po[i]
        set_cell_background(po_hdrs[i], "F1F5F9")
        set_cell_margins(po_hdrs[i], 80, 80, 80, 80)
        p = po_hdrs[i].paragraphs[0]
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(9.5)
        p.runs[0].font.name = 'Times New Roman'

    for row in po_data:
        row_cells = table_po.add_row().cells
        for col_idx, text in enumerate(row):
            row_cells[col_idx].text = text
            row_cells[col_idx].width = col_widths_po[col_idx]
            set_cell_margins(row_cells[col_idx], 60, 60, 60, 60)
            p = row_cells[col_idx].paragraphs[0]
            p.runs[0].font.size = Pt(8.5)
            p.runs[0].font.name = 'Times New Roman'
            if col_idx in [0, 2]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].font.bold = True

    doc.add_page_break()

    # -------------------------------------------------------------
    # PAGE 2: PSO JUSTIFICATION TABLE & ABSTRACT
    # -------------------------------------------------------------
    add_sub_header("PSO JUSTIFICATION")

    pso_data = [
        ("PSO1", "Apply Computing Knowledge", "3", "Applies core computing, computer networking, and cybersecurity principles to monitor Layer-2 Address Resolution Protocol (RFC 826) traffic, evaluate dynamic IP-MAC bindings, and detect rogue spoofing anomalies."),
        ("PSO2", "Problem-Solving Skills", "3", "Solves complex local network security challenges by designing multi-tier algorithmic detectors (Dynamic ARP Inspection, Sliding-Window rate limiters, MAC flip-flop trackers) and automated Gratuitous ARP healing engines."),
        ("PSO3", "Software Development", "3", "Develops and containerizes a production-grade full-stack monitoring application using Java 17 (Spring Boot 3) backend, Docker containerization, real-time WebSocket event streaming, and a high-performance React 18 frontend dashboard.")
    ]

    table_pso = doc.add_table(rows=1, cols=4)
    table_pso.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_pso.autofit = False

    pso_hdrs = table_pso.rows[0].cells
    for i, h in enumerate(po_cols):
        pso_hdrs[i].text = h.replace("PO", "PSO")
        pso_hdrs[i].width = col_widths_po[i]
        set_cell_background(pso_hdrs[i], "F1F5F9")
        set_cell_margins(pso_hdrs[i], 80, 80, 80, 80)
        p = pso_hdrs[i].paragraphs[0]
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(9.5)
        p.runs[0].font.name = 'Times New Roman'

    for row in pso_data:
        row_cells = table_pso.add_row().cells
        for col_idx, text in enumerate(row):
            row_cells[col_idx].text = text
            row_cells[col_idx].width = col_widths_po[col_idx]
            set_cell_margins(row_cells[col_idx], 60, 60, 60, 60)
            p = row_cells[col_idx].paragraphs[0]
            p.runs[0].font.size = Pt(8.5)
            p.runs[0].font.name = 'Times New Roman'
            if col_idx in [0, 2]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].font.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # ABSTRACT
    add_section_header("ABSTRACT")
    add_body_p(
        "The Address Resolution Protocol (ARP - RFC 826) is an indispensable protocol in IPv4 local area networks responsible for resolving Layer-3 IP addresses into Layer-2 Ethernet MAC addresses. However, due to its stateless design and lack of cryptographic authentication, ARP is inherently vulnerable to cache poisoning, spoofing, and Man-in-the-Middle (MITM) attacks. Operating systems unconditionally update their internal ARP tables upon receiving unsolicited ARP replies, enabling malicious actors on the local broadcast domain to intercept, modify, or drop sensitive network communication.\n\n"
        "This project presents the ARP Network Monitoring System, a full-stack defensive platform built with Java 17 (Spring Boot 3) and React 18. The system models a virtual Class-C subnet (192.168.1.0/24) comprising a Default Gateway, DNS/Web Server, Workstation Hosts, an Attacker node, and a Layer-2 Switch with isolated in-memory ARP caches. To proactively safeguard the network, the application implements a multi-tiered anomaly detection suite comprising Dynamic ARP Inspection (DAI) baseline verification, Gateway Watchdog monitoring, a sliding-window rate limiter (>18 pps), MAC flip-flop churn analysis (oscillating transitions within 5 seconds), and malformed frame checking. Upon threat identification, the self-healing mitigation engine broadcasts authoritative Gratuitous ARP (GARP) frames to cleanse corrupted host caches and enables switch port quarantine. Docker containerization and Docker Compose orchestration ensure portability across production and testing environments. With real-time bidirectional WebSocket streaming, HTML5 Canvas hardware simulation animations, and Wireshark-level deep frame inspection, the platform serves as an effective, interactive tool for network security education, protocol analysis, and intrusion detection."
    )

    add_body_p(
        "Address Resolution Protocol (RFC 826), ARP Cache Poisoning, Man-In-The-Middle (MITM), Dynamic ARP Inspection (DAI), Layer-2 Ethernet Security, Spring Boot 3, React 18, WebSockets, Docker Containerization, Self-Healing Network, Gratuitous ARP.",
        bold_prefix="KEYWORDS: "
    )

    # -------------------------------------------------------------
    # PAGE 3: INTRODUCTION & PROJECT DESCRIPTION
    # -------------------------------------------------------------
    add_section_header("1. INTRODUCTION")
    add_body_p(
        "In modern computer networks, seamless communication between networked devices relies on the coordination of protocol layers in the OSI and TCP/IP models. At the boundary between the Network Layer (Layer 3) and Data Link Layer (Layer 2), the Address Resolution Protocol (ARP), formalized in RFC 826, serves the vital function of translating logical 32-bit IPv4 addresses into physical 48-bit Ethernet Media Access Control (MAC) addresses. Whenever a host needs to transmit an IP packet to a target on the same subnet, it must determine the target's physical MAC address via an ARP Request broadcast."
    )
    add_body_p(
        "Despite its ubiquitous deployment across global enterprise and campus LANs, ARP was conceived in an era of trusted computing environments and completely lacks authentication mechanisms. Any node on the local broadcast segment can transmit unsolicited ARP reply frames containing fraudulent IP-to-MAC associations. Target hosts, operating on stateless trust, accept these frames and overwrite their internal memory caches without verification. This vulnerability forms the basis for ARP Spoofing, Gateway Impersonation, and Man-in-the-Middle (MITM) attacks, allowing an adversary to silently harvest unencrypted credentials, hijack sessions, or execute Denial-of-Service (DoS) floods."
    )
    add_body_p(
        "Traditional detection mechanisms often rely on static table bindings or expensive proprietary hardware switches. This project bridges this gap by delivering a lightweight, proactive, web-based ARP monitoring and self-healing laboratory. By pairing a robust Java Spring Boot state machine with a high-performance React dashboard, the system provides transparent visibility into Layer-2 protocol operations, detects unauthorized mapping changes with low latency, and dynamically restores network integrity."
    )

    add_section_header("2. PROJECT DESCRIPTION")
    add_body_p(
        "The primary goal of the ARP Network Monitoring System is to simulate, detect, analyze, and mitigate abnormal changes in Layer-2 ARP mappings within a local subnet. The project establishes an isolated virtual broadcast environment that mirrors realistic enterprise network topologies while eliminating the risk of accidental disruption to physical production networks."
    )

    add_sub_header("2.1 Objectives")
    add_body_p("1. Model a realistic virtual Ethernet subnet (192.168.1.0/24) with independent host ARP caches, TTL aging timers, and packet counters.")
    add_body_p("2. Implement multi-vector attack simulations including bidirectional MITM poisoning, Gateway hijacking, Gratuitous ARP storm flooding, and MAC flip-flop churn.")
    add_body_p("3. Develop a multi-tier algorithmic detection engine combining Dynamic ARP Inspection (DAI), Gateway Watchdog, sliding-window rate tracking, and churn heuristics.")
    add_body_p("4. Implement active automated mitigation via authoritative Gratuitous ARP broadcast healing and switch port quarantine.")
    add_body_p("5. Deliver a responsive, high-contrast visual dashboard featuring 60FPS packet animations, Wireshark-Lite deep frame decoding, and forensic reporting.")

    doc.add_page_break()

    # -------------------------------------------------------------
    # PAGE 4: METHODOLOGY & SYSTEM ARCHITECTURE
    # -------------------------------------------------------------
    add_section_header("3. METHODOLOGY & PROPOSED SYSTEM")
    add_body_p(
        "The proposed system employs a modular, pipeline-based methodology encompassing traffic generation, protocol decoding, multi-tiered heuristic analysis, and automated defense."
    )

    add_sub_header("3.1 Multi-Tier Anomaly Detection Algorithms")
    add_body_p("The detection engine (ARPDetector.java) evaluates every Layer-2 Ethernet and ARP frame traversing the virtual switch through five complementary inspection stages:")
    add_body_p("• Dynamic ARP Inspection (DAI) Baseline Verification: Compares incoming Sender IP and Sender MAC tuples (SPA, SHA) against a trusted DHCP Snooping and static binding hash map. Any unauthorized association triggers an immediate DAI_MISMATCH anomaly alert.")
    add_body_p("• Critical Gateway Watchdog: Maintains high-priority tracking of the Default Gateway (192.168.1.1) and DNS Server (192.168.1.10). Any frame claiming the Gateway IP with a non-baseline MAC instantly escalates the composite Threat Score to CRITICAL (85-100%).")
    add_body_p("• Sliding-Window Rate Limiter & Flood Detector: Tracks packet frequency in a rolling 2000ms window. Sustained bursts exceeding 18 total packets/sec or >6 Gratuitous ARP packets/sec trigger a DoS ARP_STORM alert.")
    add_body_p("• MAC Flip-Flop Churn Tracker: Maintains a FIFO queue of recent MAC transitions per IP address. Rapid oscillating transitions (e.g., MAC_A -> MAC_B -> MAC_A within 5 seconds) identify race-condition poisoning.")
    add_body_p("• Malformed Frame & Bogon Checker: Inspects protocol fields for RFC compliance, flagging broadcast sender MACs (FF:FF:FF:FF:FF:FF in SHA), all-zero MACs (00:00:00:00:00:00), or multicast source bindings.")

    add_sub_header("3.2 Automated Defense & Self-Healing Engine")
    add_body_p("Upon detecting an anomaly, the mitigation suite (MitigationEngine.java) executes automated remediation:")
    add_body_p("1. Authoritative Gratuitous ARP Healing: The engine constructs genuine ARP Reply frames containing the authoritative MAC for the victim IP and floods them across the subnet, immediately overwriting poisoned cache entries.")
    add_body_p("2. Layer-2 Switch Port Quarantine: Adds the offending attacker MAC to the virtual switch's drop table, actively isolating the rogue port from forwarding frames.")
    add_body_p("3. Static Binding Script Deployment: Generates OS-specific scripts for Windows (netsh), Linux (ip neigh), and Cisco IOS (ip arp inspection vlan) to permanently lock static mappings.")

    add_section_header("4. SYSTEM ARCHITECTURE")
    add_body_p(
        "The architecture is structured across four decoupled layers ensuring high scalability, concurrency, and real-time responsiveness:"
    )

    # Insert Architecture Diagram Image
    img_path = "c:\\Users\\divib\\OneDrive\\Desktop\\CN Project\\system_architecture.jpg"
    if os.path.exists(img_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(8)
        p_img.paragraph_format.space_after = Pt(6)
        run_img = p_img.add_run()
        run_img.add_picture(img_path, width=Inches(6.2))
        
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(10)
        r_cap = p_cap.add_run("Figure 4.1: System Architecture of the ARP Network Monitoring & Defense System")
        r_cap.font.size = Pt(9.5)
        r_cap.font.name = 'Times New Roman'
        r_cap.bold = True

    add_body_p("1. Frontend Layer (React 18 + HTML5 Canvas): Renders interactive topology hardware graphics (Router antennas, Server racks, Workstation LCDs, Attacker laptop), 60FPS packet trails, Wireshark tables, and telemetry charts.")
    add_body_p("2. Communication Layer (Spring WebSockets + REST): Provides bi-directional, thread-safe event streaming via `/ws/traffic` using ConcurrentWebSocketSessionDecorator, pushing real-time packet payloads and security alerts.")
    add_body_p("3. Business & Simulation Layer (Spring Boot 3 Core): Manages virtual host state machines, background traffic generation (@Scheduled loops), attack injections, detection heuristics, and self-healing algorithms.")
    add_body_p("4. Persistence Layer (In-Memory Database Service): Maintains rolling buffers of recent Layer-2 frames, anomaly audit logs, and hardware device binding states.")

    doc.add_page_break()

    # -------------------------------------------------------------
    # PAGE 5: IMPLEMENTATION DETAILS & CODE MANIFEST
    # -------------------------------------------------------------
    add_section_header("5. IMPLEMENTATION DETAILS")
    add_body_p(
        "The project is implemented using Java 17 and Spring Boot 3.3.3 on the backend, paired with a React 18 frontend packaged for zero-dependency standalone execution."
    )

    add_sub_header("5.1 Key Backend Source Files")
    backend_files = [
        ("ArpMonitorApplication.java", "Main Spring Boot bootstrap class with @EnableScheduling background traffic loops and banner initialization."),
        ("SimulationEngine.java", "Core virtual subnet state machine managing Router, Server, Hosts A/B/C, Kali node, and individual in-memory ARP caches."),
        ("ARPDetector.java", "Multi-tier anomaly detection algorithms (DAI verification, Gateway Watchdog, sliding-window rate limiting, and MAC churn heuristics)."),
        ("MitigationEngine.java", "Self-healing engine broadcasting authoritative Gratuitous ARP frames, managing switch port isolation, and generating OS static binding scripts."),
        ("AttackSimulator.java", "Packet crafting engine executing MITM poisoning, Gateway spoofing, GARP storms, and custom Layer-2 frame injection."),
        ("TrafficWebSocketHandler.java", "Thread-safe WebSocket broadcaster utilizing ConcurrentWebSocketSessionDecorator for concurrent frame streaming."),
        ("ARPPacket.java & NetworkNode.java", "Domain entity models defining Ethernet II/ARP frame fields, hardware specs, and in-memory cache structures.")
    ]

    table_bf = doc.add_table(rows=1, cols=2)
    table_bf.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_bf.autofit = False
    table_bf.rows[0].cells[0].text = "Source File"
    table_bf.rows[0].cells[0].width = Inches(2.2)
    table_bf.rows[0].cells[1].text = "Functional Responsibility"
    table_bf.rows[0].cells[1].width = Inches(4.3)
    set_cell_background(table_bf.rows[0].cells[0], "F1F5F9")
    set_cell_background(table_bf.rows[0].cells[1], "F1F5F9")
    table_bf.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
    table_bf.rows[0].cells[1].paragraphs[0].runs[0].font.bold = True

    for fname, fdesc in backend_files:
        row_cells = table_bf.add_row().cells
        row_cells[0].text = fname
        row_cells[0].width = Inches(2.2)
        row_cells[1].text = fdesc
        row_cells[1].width = Inches(4.3)
        set_cell_margins(row_cells[0], 50, 50, 50, 50)
        set_cell_margins(row_cells[1], 50, 50, 50, 50)
        row_cells[0].paragraphs[0].runs[0].font.size = Pt(8.5)
        row_cells[0].paragraphs[0].runs[0].font.name = 'Consolas'
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[1].paragraphs[0].runs[0].font.size = Pt(8.5)
        row_cells[1].paragraphs[0].runs[0].font.name = 'Times New Roman'

    add_sub_header("5.2 Key Frontend Source Files")
    frontend_files = [
        ("static/index.html", "Unified Single-Page SOC Dashboard integrating React 18, Tailwind CSS, Chart.js, and HTML5 Canvas 60FPS renderer."),
        ("TopologyCanvas.jsx", "Interactive network topology visualizer with animated device graphics, link cables, packet pulse beams, and selection halos."),
        ("ArpCacheMatrix.jsx", "Synchronized grid displaying side-by-side internal ARP tables with highlighted poisoned entries."),
        ("WiresharkAnalyzer.jsx", "Streaming Layer-2 packet analyzer table with Opcode badges (REQ, REP, GARP) and deep packet inspection modals."),
        ("AlertsFeed.jsx", "Real-time security incident feed displaying alert severity, victim IP, claimed rogue MAC, and mitigation steps.")
    ]

    table_ff = doc.add_table(rows=1, cols=2)
    table_ff.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_ff.autofit = False
    table_ff.rows[0].cells[0].text = "Frontend Component"
    table_ff.rows[0].cells[0].width = Inches(2.2)
    table_ff.rows[0].cells[1].text = "Functional Responsibility"
    table_ff.rows[0].cells[1].width = Inches(4.3)
    set_cell_background(table_ff.rows[0].cells[0], "F1F5F9")
    set_cell_background(table_ff.rows[0].cells[1], "F1F5F9")
    table_ff.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
    table_ff.rows[0].cells[1].paragraphs[0].runs[0].font.bold = True

    for fname, fdesc in frontend_files:
        row_cells = table_ff.add_row().cells
        row_cells[0].text = fname
        row_cells[0].width = Inches(2.2)
        row_cells[1].text = fdesc
        row_cells[1].width = Inches(4.3)
        set_cell_margins(row_cells[0], 50, 50, 50, 50)
        set_cell_margins(row_cells[1], 50, 50, 50, 50)
        row_cells[0].paragraphs[0].runs[0].font.size = Pt(8.5)
        row_cells[0].paragraphs[0].runs[0].font.name = 'Consolas'
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[1].paragraphs[0].runs[0].font.size = Pt(8.5)
        row_cells[1].paragraphs[0].runs[0].font.name = 'Times New Roman'

    doc.add_page_break()

    # -------------------------------------------------------------
    # PAGE 6: RESULTS, DISCUSSION & EVALUATION
    # -------------------------------------------------------------
    add_section_header("6. RESULTS AND DISCUSSIONS")
    add_body_p(
        "Comprehensive experimental testing was conducted across multiple scenarios to validate the detection speed, accuracy, and self-healing effectiveness of the system."
    )

    add_sub_header("6.1 Experimental Scenarios & Observations")
    add_body_p("• Scenario 1: Baseline Network Traffic: Under normal operating conditions, background ARP requests and replies traverse the virtual switch every 1.5 seconds. The DAI inspection engine continuously verifies that Host A (192.168.1.101) maps to 00:1A:2B:3C:4D:01 and Gateway (192.168.1.1) maps to 00:1A:2B:3C:4D:GW. Threat score remains 0% (NORMAL).")
    add_body_p("• Scenario 2: MITM Poisoning Attack Injection: Triggering a bidirectional MITM attack injects forged ARP reply frames associating the Gateway IP with the Attacker MAC (AA:BB:CC:DD:EE:66). Host A's internal cache is immediately poisoned, turning its LCD monitor red with a '⚠️ POISON' alert. The DAI detector flags the mismatch in under 2ms, generating a CRITICAL incident log.")
    add_body_p("• Scenario 3: Gateway Impersonation: Forged Gratuitous ARP frames claiming Gateway ownership cause the Gateway Watchdog to immediately escalate the Threat Score to 100%. An audio alert sounds, and the affected node hales red on the 60FPS canvas.")
    add_body_p("• Scenario 4: Automated Self-Healing & Cleanse: Clicking 'Broadcast Heal' (or enabling 'Auto-Mitigate') triggers the mitigation engine to flood authoritative Gratuitous ARP packets containing legitimate MACs. Host A's internal cache is instantly overwritten and restored, returning the threat status to GREEN (0%).")

    add_sub_header("6.2 Security Feature Comparison Table")
    comp_headers = ["Security Mechanism", "Detection Speed", "MITM Protection", "Self-Healing", "Operational Overhead"]
    comp_rows = [
        ["Standard ARP (RFC 826)", "None (Vulnerable)", "No", "No", "Low (Unauthenticated)"],
        ["Static ARP Bindings", "Static Only", "Partial", "Manual Reset Only", "High (Manual Maintenance)"],
        ["Hardware Switch DAI", "Hardware Latency", "Yes (Drops frames)", "No (No Healing)", "High (Enterprise Switch Cost)"],
        ["Proposed ARP SOC System", "< 2ms (Real-time)", "Yes (DAI + Churn)", "Yes (Auto GARP Heal)", "Low (Automated Full-Stack)"]
    ]

    table_comp = doc.add_table(rows=1, cols=5)
    table_comp.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_comp.autofit = False

    comp_widths = [Inches(1.8), Inches(1.2), Inches(1.1), Inches(1.1), Inches(1.3)]
    for i, h in enumerate(comp_headers):
        cell = table_comp.rows[0].cells[i]
        cell.text = h
        cell.width = comp_widths[i]
        set_cell_background(cell, "F1F5F9")
        set_cell_margins(cell, 60, 60, 50, 50)
        p = cell.paragraphs[0]
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(8.5)
        p.runs[0].font.name = 'Times New Roman'

    for row in comp_rows:
        row_cells = table_comp.add_row().cells
        for col_idx, text in enumerate(row):
            row_cells[col_idx].text = text
            row_cells[col_idx].width = comp_widths[col_idx]
            set_cell_margins(row_cells[col_idx], 50, 50, 50, 50)
            p = row_cells[col_idx].paragraphs[0]
            p.runs[0].font.size = Pt(8)
            p.runs[0].font.name = 'Times New Roman'
            if col_idx == 0:
                p.runs[0].font.bold = True

    # -------------------------------------------------------------
    # PAGE 7: CONCLUSION & FUTURE ENHANCEMENTS
    # -------------------------------------------------------------
    add_section_header("7. CONCLUSION")
    add_body_p(
        "The ARP Network Monitoring System successfully addresses the fundamental vulnerabilities inherent in the Address Resolution Protocol (RFC 826). By combining a high-performance Java Spring Boot 3 backend state machine with an interactive React 18 operations dashboard, the project demonstrates how automated, multi-tiered anomaly detection can eliminate the risks of ARP cache poisoning, Man-In-The-Middle attacks, and subnet hijacking. The integration of real-time WebSocket telemetry, HTML5 Canvas 60FPS packet animation, and authoritative Gratuitous ARP healing establishes a practical, production-ready framework for network security education and active defense monitoring."
    )

    add_section_header("8. FUTURE ENHANCEMENTS")
    add_body_p("• Machine Learning Behavioral Profiling: Integrating LSTM recurrent neural networks or Isolation Forests to profile subtle, low-frequency cache poisoning patterns without static thresholding.")
    add_body_p("• Physical Raw Socket Capture (libpcap / Scapy): Extending the backend packet ingestion pipeline to capture live physical Ethernet frames on mirror/SPAN switch ports.")
    add_body_p("• IPv6 Neighbor Discovery Protocol (NDP) Monitoring: Expanding detection support to IPv6 ICMPv6 Neighbor Solicitation and Advertisement spoofing, incorporating Secure Neighbor Discovery (SEND - RFC 3971) verification.")
    add_body_p("• Enterprise SIEM Integration: Enabling automated Syslog / CEF alert forwarding to enterprise Security Operations Center (SOC) platforms such as Splunk, Elastic SIEM, or Wazuh.")

    add_section_header("9. REFERENCES")
    add_body_p("1. Plummer, D. C. (1982). \"An Ethernet Address Resolution Protocol: Converting Network Protocol Addresses to 48.bit Ethernet Address for Hardware Transmission on Ethernet Hardware,\" RFC 826, Internet Engineering Task Force (IETF).")
    add_body_p("2. Cheshire, S. (2008). \"IPv4 Address Conflict Detection,\" RFC 5227, Internet Engineering Task Force (IETF).")
    add_body_p("3. Arkko, J., et al. (2005). \"SEcure Neighbor Discovery (SEND),\" RFC 3971, Internet Engineering Task Force (IETF).")
    add_body_p("4. Tanenbaum, A. S., & Wetherall, D. J. (2011). Computer Networks (5th ed.). Prentice Hall.")
    add_body_p("5. Kurose, J. F., & Ross, K. W. (2021). Computer Networking: A Top-Down Approach (8th ed.). Pearson.")

    # Save document with fallback
    output_path = "c:\\Users\\divib\\OneDrive\\Desktop\\CN Project\\ARP_Network_Monitoring_System_Mini_Project_Report.docx"
    output_fallback = "c:\\Users\\divib\\OneDrive\\Desktop\\CN Project\\ARP_Network_Monitoring_System_Report_With_Architecture.docx"
    
    try:
        doc.save(output_path)
        print(f"[+] Word Document Report successfully saved at: {output_path}")
    except PermissionError:
        doc.save(output_fallback)
        print(f"[!] Primary file was open in Word. Saved updated report with diagram to: {output_fallback}")

if __name__ == "__main__":
    create_report()
